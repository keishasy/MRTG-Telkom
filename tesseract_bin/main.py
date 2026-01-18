import pandas as pd
import os
import requests
from bs4 import BeautifulSoup
import re
import time
import tkinter as tk
from tkinter import Label, Button, Frame
from PIL import Image, ImageTk
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import xlsxwriter
from datetime import datetime
import calendar

# --- KONFIGURASI ---
FILE_INPUT = 'input_data.xlsx'
LOGO_PATH = 'assets/telkom.png'
FOLDER_GAMBAR = 'downloaded_images'

# DAFTAR URL TARGET (Robot akan cek satu per satu)
TARGET_URLS = [
    "http://10.62.8.136",
    "http://10.62.8.135/cacti"
]

# --- A. MODUL GUI (POPUP PEMILIH GAMBAR) ---
class ImageSelector:
    def __init__(self, images_data, sid_title):
        self.selected_image = None
        self.root = tk.Tk()
        self.root.title(f"KONFLIK DATA: {sid_title}")
        self.root.attributes('-topmost', True)
        
        # Hitung lebar jendela berdasarkan jumlah gambar
        window_width = max(600, len(images_data) * 320)
        self.root.geometry(f"{window_width}x450")

        lbl_instruksi = Label(self.root, 
                              text=f"Ditemukan {len(images_data)} grafik untuk SID: {sid_title}.\nCek Server asal & ID, lalu KLIK gambar yang benar.", 
                              font=("Arial", 11, "bold"))
        lbl_instruksi.pack(pady=10)

        # Scrollable Frame (biar aman kalau gambarnya banyak banget)
        canvas = tk.Canvas(self.root)
        frame_imgs = Frame(canvas)
        scrollbar = tk.Scrollbar(self.root, orient="horizontal", command=canvas.xview)
        canvas.configure(xscrollcommand=scrollbar.set)
        
        scrollbar.pack(side="bottom", fill="x")
        canvas.pack(side="top", fill="both", expand=True)
        canvas.create_window((0, 0), window=frame_imgs, anchor="nw")
        
        frame_imgs.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        self.photos = [] 

        for idx, (img_bytes, graph_id, source_url) in enumerate(images_data):
            try:
                img = Image.open(io.BytesIO(img_bytes))
                img.thumbnail((300, 220)) # Resize thumbnail
                photo = ImageTk.PhotoImage(img)
                self.photos.append(photo)

                col_frame = Frame(frame_imgs, borderwidth=2, relief="groove", padx=5, pady=5)
                col_frame.pack(side="left", padx=10, pady=10)

                # Info Server
                lbl_srv = Label(col_frame, text=f"Sumber: {source_url}", fg="blue", font=("Arial", 8, "bold"))
                lbl_srv.pack()

                # Tombol Gambar
                btn = Button(col_frame, image=photo, command=lambda x=img_bytes: self.select(x))
                btn.pack(pady=5)
                
                # Info ID
                Label(col_frame, text=f"ID: {graph_id}", font=("Arial", 8)).pack()
            except:
                pass

    def select(self, img_bytes):
        self.selected_image = img_bytes
        self.root.destroy()

    def show(self):
        self.root.mainloop()
        return self.selected_image

# --- B. MODUL ROBOT SCRAPER (MULTI-SERVER) ---
def download_graph_mrtg_multi(sid, tgl_start, tgl_end):
    session = requests.Session()
    all_candidates = [] # Menampung hasil dari SEMUA server
    
    # LOOPING KE SETIAP URL TARGET
    for base_url in TARGET_URLS:
        # print(f"   🔎 Cek Server: {base_url} ...")
        search_url = f"{base_url}/graph_view.php"
        
        params = {
            'action': 'preview',
            'filter': sid,
            'date1': tgl_start,
            'date2': tgl_end,
            'predefined_timespan': '0'
        }

        try:
            # 1. SEARCH
            response = session.get(search_url, params=params, timeout=5) # Timeout 5 detik per server
            if response.status_code != 200: continue

            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 2. CARI ID
            found_ids = set()
            
            # Pola Link
            links = soup.find_all('a', href=re.compile(r'local_graph_id=\d+'))
            for l in links:
                match = re.search(r'local_graph_id=(\d+)', l['href'])
                if match: found_ids.add(match.group(1))
                
            # Pola Image (Backup)
            imgs = soup.find_all('img', src=re.compile(r'local_graph_id=\d+'))
            for i in imgs:
                match = re.search(r'local_graph_id=(\d+)', i['src'])
                if match: found_ids.add(match.group(1))
            
            # 3. DOWNLOAD KANDIDAT DARI SERVER INI
            for gid in found_ids:
                img_url = f"{base_url}/graph_image.php?action=view&local_graph_id={gid}&rra_id=3"
                res = session.get(img_url, timeout=10)
                if res.status_code == 200:
                    # Simpan: (Bytes Gambar, ID, URL Asal)
                    # Kita simpan URL asal biar user tau ini gambar dari server mana
                    server_name = "136 (Utama)" if "136" in base_url else "135 (Cacti)"
                    all_candidates.append((res.content, gid, server_name))
                    
        except requests.exceptions.RequestException:
            # print(f"   ⚠️ Server {base_url} unreachable/timeout.")
            continue # Lanjut ke server berikutnya
        except Exception:
            continue

    # --- LOGIKA PENENTUAN ---
    
    if not all_candidates:
        return None # ZONK di kedua server

    final_image_bytes = None
    
    if len(all_candidates) == 1:
        # HORE! Cuma ketemu 1 gambar (bersih)
        final_image_bytes = all_candidates[0][0]
    else:
        # GANDA! Ketemu di 136 DAN 135, atau ID double
        print(f"   ⚠️  Ditemukan {len(all_candidates)} opsi grafik (Cek Popup)...")
        selector = ImageSelector(all_candidates, sid)
        final_image_bytes = selector.show()
        
        if final_image_bytes:
            print("   ✅ Gambar dipilih user.")
        else:
            print("   ❌ User skip.")
            return None

    # SIMPAN HASIL AKHIR
    if final_image_bytes:
        if not os.path.exists(FOLDER_GAMBAR): os.makedirs(FOLDER_GAMBAR)
        clean_sid = "".join([c for c in sid if c.isalnum() or c in (' ','-','_')]).strip()
        filename = f"{FOLDER_GAMBAR}/{clean_sid}.png"
        
        with open(filename, 'wb') as f:
            f.write(final_image_bytes)
        return filename
        
    return None

# --- C. FUNGSI HITUNG TANGGAL ---
def get_filter_parameters(bulan, tahun):
    start_date = f"{tahun}-{bulan:02d}-01 00:01"
    last_day = calendar.monthrange(int(tahun), bulan)[1]
    end_date = f"{tahun}-{bulan:02d}-{last_day} 23:59"
    return start_date, end_date

# --- D. INPUT USER ---
def get_user_input():
    print("\n--- KONFIGURASI ROBOT MRTG ---")
    while True:
        try:
            bulan_input = int(input("Masukkan Bulan (1-12): "))
            if 1 <= bulan_input <= 12: break
        except: pass
    while True:
        try:
            tahun_input = int(input("Masukkan Tahun (contoh 2025): "))
            if 2000 <= tahun_input <= 2100: break
        except: pass
    
    nama_bulan_indo = ["", "JANUARI", "FEBRUARI", "MARET", "APRIL", "MEI", "JUNI", 
                       "JULI", "AGUSTUS", "SEPTEMBER", "OKTOBER", "NOVEMBER", "DESEMBER"]
    return nama_bulan_indo[bulan_input], str(tahun_input), bulan_input

# --- E. LAYOUT WORD ---
def set_column_layout(section):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        docGrid = sectPr.find(qn('w:docGrid'))
        if docGrid is not None: docGrid.addprevious(cols)
        else: sectPr.append(cols)
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')
    cols.set(qn('w:equalWidth'), '1')

# --- MAIN ---
def main():
    print("=== ROBOT MRTG MULTI-SERVER (136 & 135) ===")
    print(f"Server Target: {TARGET_URLS}")
    
    # SETUP AWAL
    nama_bulan, tahun_str, bulan_angka = get_user_input()
    tgl_start, tgl_end = get_filter_parameters(bulan_angka, int(tahun_str))
    
    file_word = f"MRTG {nama_bulan} {tahun_str}.docx"
    file_excel = f"MRTG {nama_bulan} {tahun_str}.xlsx"
    
    # BACA EXCEL
    if not os.path.exists(FILE_INPUT):
        print(f"❌ File {FILE_INPUT} tidak ada.")
        return
    df = pd.read_excel(FILE_INPUT)
    df.columns = [c.lower().strip() for c in df.columns]

    # SETUP WORD
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(7)
    
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_column_layout(section)

    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(LOGO_PATH, width=Inches(1.2))
    
    p = doc.add_paragraph()
    run = p.add_run('TRAFIK MRTG BANK JATIM')
    run.font.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph(f'{len(df)} LINK TELKOM {nama_bulan} {tahun_str}').paragraph_format.space_after = Pt(2)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").paragraph_format.space_after = Pt(12)

    # SETUP EXCEL
    wb = xlsxwriter.Workbook(file_excel)
    ws = wb.add_worksheet("Rekap")
    fmt_head = wb.add_format({'bold':True, 'bg_color':'#D7E4BC', 'border':1, 'align':'center'})
    fmt_cell = wb.add_format({'border':1})
    
    headers = ["No", "Alamat + SID", "Inbound", "Outbound", "Total AVG", "Total Trafik"]
    for c, h in enumerate(headers): ws.write(0, c, h, fmt_head)
    ws.set_column('B:B', 50)

    # LOOPING
    print("🚀 Mulai Berburu Grafik...")
    
    for idx, row in df.iterrows():
        no = idx + 1
        alamat = str(row['alamat'])
        sid = str(row['sid'])
        bw = str(row['bandwidth'])
        
        teks_id = f"{no}. {alamat} SID.{sid} BW.{bw}"
        print(f"[{no}/{len(df)}] {sid}...", end=" ")
        
        # --- PANGGIL FUNGSI MULTI-SERVER ---
        gambar_path = download_graph_mrtg_multi(sid, tgl_start, tgl_end)
        
        # WORD
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        run = p.add_run(teks_id)
        run.bold = True
        run.font.size = Pt(7)
        
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_img.paragraph_format.space_after = Pt(8)
        
        if gambar_path:
            print("✅ OK")
            try:
                p_img.add_run().add_picture(gambar_path, width=Inches(2.5))
            except:
                p_img.add_run("[ERROR GAMBAR RUSAK]").font.color.rgb = RGBColor(255,0,0)
        else:
            print("❌ ZONK (Semua Server)")
            run_err = p_img.add_run("[GRAFIK TIDAK DITEMUKAN]")
            run_err.font.color.rgb = RGBColor(255, 0, 0)
            run_err.bold = True

        # EXCEL
        teks_excel = f"{alamat} SID.{sid} BW.{bw}"
        ws.write(idx+1, 0, no, fmt_cell)
        ws.write(idx+1, 1, teks_excel, fmt_cell)
        for c in range(2, 6): ws.write(idx+1, c, 0, fmt_cell)

    try:
        doc.save(file_word)
        wb.close()
        print(f"\n🎉 SELESAI! Cek file: {file_word}")
    except Exception as e:
        print(f"❌ Gagal simpan: {e}")

if __name__ == "__main__":
    main()