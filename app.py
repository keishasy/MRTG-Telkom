import streamlit as st
import pandas as pd
import calendar
from datetime import datetime
import os
import requests
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pytesseract
from PIL import Image
from openpyxl import load_workbook

# UPDATE IMPORT: Tambahkan 'Alignment'
from openpyxl.styles import PatternFill, Alignment
import base64
import sys
from pathlib import Path
from datetime import date, timedelta


def resource_path(rel_path: str) -> str:
    base = Path(getattr(sys, "_MEIPASS", Path(__file__).parent))
    return str(base / rel_path)


st.set_page_config(page_title="Dashboard Monitoring", page_icon="📡", layout="wide")

# TESSERACT
pytesseract.pytesseract.tesseract_cmd = resource_path("tesseract_bin/tesseract.exe")
os.environ["TESSDATA_PREFIX"] = resource_path("tesseract_bin/tessdata")


# FUNGSI GAMBAR
def img_to_base64(rel_path: str) -> str:
    p = resource_path(rel_path)
    if not os.path.exists(p):
        return ""
    with open(p, "rb") as f:
        return base64.b64encode(f.read()).decode()


def asset_exists(rel_path: str) -> bool:
    return os.path.exists(resource_path(rel_path))


# LOAD ICON GLOBALLY
loading_icon_b64 = img_to_base64("assets/loading.png")
add_icon_b64 = img_to_base64("assets/add.png")

# CSS GLOBAL
st.markdown(
    f"""
    <style>
        #MainMenu {{ 
            visibility: hidden; 
        }}

        header {{ 
            visibility: hidden; 
        }}

        footer {{ 
            visibility: hidden; 
        }}

        .red-strip {{ 
            position: fixed; 
            top: 0; 
            left: 0; 
            width: 15px; 
            height: 100vh; 
            background-color: #EE2D24; 
            z-index: 9999; 
        }}
        
        .block-container {{ 
            padding-top: 1rem !important; 
            padding-left: 3.5rem !important; 
        }}

        .stProgress > div > div > div > div {{ 
            background-color: #EE2D24; 
        }}

        [data-testid="stBorderWrapper"] {{ 
            border: 2px solid #000 !important; 
            border-radius: 15px !important; 
            padding: 30px !important; 
            background-color: white; 
        }}

        .main-title {{ 
            font-size: 28px; 
            font-weight: 800; 
            color: #1a1a1a; 
            margin-bottom: 5px; 
        }}

        .sub-title {{ 
            font-size: 14px; 
            color: #555; 
            margin-bottom: 30px; 
        }}

        .upload-note {{ 
            font-size: 12px; 
            color: #d32f2f; 
            font-style: italic; 
            margin-top: -10px; 
            margin-bottom: 20px; 
        }}
        
        .card-success {{ 
            background-color: #ECFDF3; 
            border-left: 6px solid #16A34A; 
            border-radius: 14px; 
            padding: 16px 18px; 
            margin-bottom: 14px; 
        }}

        .card-failed {{ 
            background-color: #FEF2F2; 
            border-left: 6px solid #DC2626; 
            border-radius: 14px; 
            padding: 16px 18px; 
            margin-bottom: 14px; 
        }}

        .card-multiple {{ 
            background-color: #FFFBEB; 
            border-left: 6px solid #F59E0B; 
            border-radius: 14px; 
            padding: 16px 18px; 
            margin-bottom: 14px; 
        }}

        .card-title {{ 
            font-weight: 600; 
            font-size: 14px; 
            margin-bottom: 4px; 
        }}

        .card-meta {{ 
            font-size: 12px; 
            color: #555; 
        }}
        
        .metric-box {{ 
            background-color: white; 
            border: 1px solid #e5e7eb; 
            border-radius: 12px; 
            padding: 15px 10px; 
            box-shadow: 0 2px 4px rgba(0,0,0,0.03); 
            text-align: center; 
            display: flex; 
            flex-direction: column; 
            align-items: center; 
            justify-content: center; 
            height: 100%; 
        }}

        .metric-title {{ 
            font-size: 11px; 
            font-weight: 700; 
            color: #6b7280; 
            text-transform: uppercase; 
            letter-spacing: 0.8px; 
            margin-bottom: 8px; 
        }}

        .metric-value {{ 
            font-size: 32px; 
            font-weight: 800; 
            color: #111827; 
            line-height: 1; 
        }}

        .mb-blue {{ 
            border-bottom: 4px solid #3b82f6; 
        }}

        .mb-green {{ 
            border-bottom: 4px solid #22c55e; 
        }}

        .mb-yellow {{ 
            border-bottom: 4px solid #eab308; 
        }}

        .mb-red {{ 
            border-bottom: 4px solid #ef4444; 
        }}
        
        div[role="radiogroup"] {{ 
            justify-content: center !important; 
        }}
        
        [data-testid="stFileUploader"] button {{ 
            background-color: #EE2D24 !important; 
            border: 2px solid #EE2D24 !important; 
            color: white !important; 
            border-radius: 50px !important; 
        }}

        [data-testid="stFileUploader"] section {{ 
            background-color: #f8f9fa !important; 
            border-radius: 15px !important; 
            border: 1px dashed #ccc !important; 
        }}
        
        /* CSS TOAST ANIMATION */
        @keyframes slideIn {{ 
            from {{ 
                transform: translateX(100%); 
                opacity: 0; 
            }} 
            to {{ 
                transform: translateX(0); 
                opacity: 1; 
            }} 
        }}

        @keyframes spin {{ 
            0% {{ 
                transform: rotate(0deg); 
            }} 
            100% {{ 
                transform: rotate(360deg); 
            }} 
        }}
        
        .custom-toast {{
            position: fixed; 
            top: 20px; 
            right: 20px;
            background-color: white; 
            border-left: 5px solid #EE2D24;
            padding: 15px 20px; 
            border-radius: 8px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.15);
            display: flex; 
            align-items: center; 
            gap: 15px;
            z-index: 999999; 
            animation: slideIn 0.3s ease-out forwards;
            min-width: 320px;
        }}

        .toast-icon {{
            width: 30px; 
            height: 30px;
            background-image: url('data:image/png;base64,{loading_icon_b64}');
            background-size: contain; 
            background-repeat: no-repeat; 
            background-position: center;
            animation: spin 1s linear infinite;
        }}

        .toast-content {{ 
            font-family: 'Segoe UI', sans-serif; 
            flex-grow: 1; 
        }}

        .toast-title {{ 
            font-weight: 800; 
            font-size: 14px; 
            color: #333; 
            margin-bottom: 2px; 
        }}

        .toast-desc {{ 
            font-size: 13px; 
            color: #666; 
            font-weight: 500; 
        }}

        .toast-progress {{ 
            font-size: 11px; 
            color: #999; 
            margin-top: 2px; 
        }}
    </style>

    <div class="red-strip"></div>
    """,
    unsafe_allow_html=True,
)


# BACKEND
# SCRAPPING DATA DARI WEB GRAPH
def _daterange(start_date: date, end_date: date):
    cur = start_date
    while cur <= end_date:
        yield cur
        cur += timedelta(days=1)


def scrape_dynamic(sid, start_dt_str, end_dt_str, rra_id):
    """
    Mengambil grafik dari 3 server dengan rentang waktu SPESIFIK (Unix Timestamp).
    Ini memastikan grafik 'From... To...' sesuai input user.
    """
    servers = [
        "http://10.62.8.136/cacti",
        "http://10.62.8.135/cacti",
        "http://10.62.8.132/cacti",
    ]

    try:
        dt_start_obj = datetime.strptime(start_dt_str, "%Y-%m-%d %H:%M")
        dt_end_obj = datetime.strptime(end_dt_str, "%Y-%m-%d %H:%M")

        ts_start = int(dt_start_obj.timestamp())
        ts_end = int(dt_end_obj.timestamp())
    except:
        ts_start = 0
        ts_end = 0

    combined_graphs = []

    for base in servers:
        try:
            params = {
                "action": "preview",
                "filter": sid,
                "date1": start_dt_str,
                "date2": end_dt_str,
            }
            resp = requests.get(f"{base}/graph_view.php", params=params, timeout=10)
            graph_ids = re.findall(r"local_graph_id=(\d+)", resp.text)

            for g_id in list(set(graph_ids)):
                img_url = (
                    f"{base}/graph_image.php?"
                    f"action=view&local_graph_id={g_id}&rra_id={rra_id}"
                    f"&graph_start={ts_start}&graph_end={ts_end}"
                )

                try:
                    ip_part = base.split("//")[1].split("/")[0]
                    server_label = ip_part.split(".")[-1]
                except:
                    server_label = "Unknown"

                combined_graphs.append(
                    {
                        "url": img_url,
                        "server": server_label,
                        "date1": start_dt_str,
                        "date2": end_dt_str,
                        "rra_id": rra_id,
                    }
                )
        except Exception as e:
            continue

    return combined_graphs


# HITUNG KE KBPS
def convert_to_kbps(value_str, unit):
    try:
        s = str(value_str).strip()
        if s == "" or s.lower() in ["nan", "-nan", "none"]:
            return 0.0

        s = s.replace(",", ".")
        val = float(s)

        u = (unit or "").strip()

        if u == "k" or u == "K":
            return val

        elif u == "M":
            return val * 1024

        elif u == "m":
            return val / 1000

        elif u == "G" or u == "g":
            return val * 1_000_000

        elif u == "":
            return val / 1000

        else:
            return val / 1_000_000

    except Exception as e:
        return 0.0


# OCR
def ocr_extract_data(image_path):
    try:
        img = Image.open(image_path)
        w, h = img.size

        img = img.crop((0, int(h * 0.70), w, h))

        img = img.resize((img.width * 3, img.height * 3), Image.Resampling.LANCZOS)
        img = img.convert("L")
        img = img.point(lambda x: 0 if x < 165 else 255, "1")

        text = pytesseract.image_to_string(img, config=r"--oem 3 --psm 6")

        def clean_num(s):
            s = str(s).strip().replace(",", ".")
            s = re.sub(r"\.+", ".", s)
            return s.strip(".")

        def fmt(val, unit):
            val = clean_num(val)
            unit = (unit or "").strip()
            return f"{val} {unit}".strip()

        max_word = r"Max(?:imum|imurn|irnum)?"

        pat_in = rf"Inbound.*?Average:\s*([\d\.,]+)\s*([kKmMgG]?)\s*{max_word}:\s*([\d\.,]+)\s*([kKmMgG]?)"
        pat_out = rf"Outbound.*?Average:\s*([\d\.,]+)\s*([kKmMgG]?)\s*{max_word}:\s*([\d\.,]+)\s*([kKmMgG]?)"

        m_in = re.search(pat_in, text, re.IGNORECASE | re.DOTALL)
        m_out = re.search(pat_out, text, re.IGNORECASE | re.DOTALL)

        avg_in = fmt(m_in.group(1), m_in.group(2)) if m_in else "0"
        max_in = fmt(m_in.group(3), m_in.group(4)) if m_in else "0"
        avg_out = fmt(m_out.group(1), m_out.group(2)) if m_out else "0"
        max_out = fmt(m_out.group(3), m_out.group(4)) if m_out else "0"

        return avg_in, avg_out, max_in, max_out

    except:
        return "0", "0", "0", "0"


# GENERATE WORD
def generate_clean_word(data, target_date_str):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = (
        section.right_margin
    ) = Inches(1)
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "720")

    if asset_exists("assets/telkom.jpg"):
        logo_p = doc.add_paragraph()
        logo_p.add_run().add_picture(
            resource_path("assets/telkom.jpg"), width=Inches(1)
        )
        logo_p.alignment = 0

    title_p1 = doc.add_paragraph()
    r1 = title_p1.add_run("LAPORAN TRAFIK MRTG\nBANK JATIM")
    r1.font.bold = True
    r1.font.size = Pt(16)

    total_cabang = len(data)
    info_p = doc.add_paragraph()
    r_info = info_p.add_run(f"{total_cabang} LINK TELKOM")
    r_info.font.bold = True
    r_info.font.size = Pt(14)

    date_p = doc.add_paragraph()
    r_date = date_p.add_run(f"{target_date_str}")
    r_date.font.size = Pt(12)
    doc.add_paragraph("")

    for i, item in enumerate(data, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_together = True

        header = f"{i}. {item['alamat']} SID.{item['sid']} BW.{item['bw']}"
        r = p.add_run(header)
        r.font.size = Pt(7)
        r.font.bold = True

        if item.get("selected_url"):
            try:
                img = requests.get(item["selected_url"], timeout=12).content
                tmp_img = f"tmp_{i}.png"
                with open(tmp_img, "wb") as f:
                    f.write(img)
                img_p = doc.add_paragraph()
                img_p.paragraph_format.keep_together = True
                img_p.add_run().add_picture(tmp_img, width=Inches(2.8))
                os.remove(tmp_img)
            except:
                err = doc.add_paragraph()
                rr = err.add_run("[Tidak Ditemukan Grafik]")
                rr.font.size = Pt(6)
                rr.font.color.rgb = RGBColor(255, 0, 0)
                rr.font.italic = True
        else:
            err = doc.add_paragraph()
            rr = err.add_run("[Tidak Ditemukan Grafik]")
            rr.font.size = Pt(6)
            rr.font.color.rgb = RGBColor(255, 0, 0)
            rr.font.italic = True

        doc.add_paragraph("")

    safe_date = target_date_str.replace("/", "-")
    out_name = f"Laporan MRTG BANK JATIM {safe_date}.docx"
    doc.save(out_name)
    return out_name


# GENERATE EXCEL
def generate_excel_report(data, target_date_str, date_obj=None):
    try:
        if date_obj:
            year = date_obj.year
            month = date_obj.month
            days_in_month = calendar.monthrange(year, month)[1]
        else:
            first_date_part = target_date_str.split(" ")[0]
            dt_obj = datetime.strptime(first_date_part, "%Y-%m-%d")
            year = dt_obj.year
            month = dt_obj.month
            days_in_month = calendar.monthrange(year, month)[1]
    except:
        days_in_month = 30

    rows = []

    for i, item in enumerate(data, start=1):
        raw_avg_in = item.get("avg_in", "0")
        raw_avg_out = item.get("avg_out", "0")

        def split_val_unit(s):
            s = str(s).strip()
            match = re.match(r"([\d\.,]+)\s*([a-zA-Z]*)", s)
            if match:
                return match.group(1), match.group(2)
            return s, ""

        val_in, unit_in = split_val_unit(raw_avg_in)
        val_out, unit_out = split_val_unit(raw_avg_out)

        in_kbps = convert_to_kbps(val_in, unit_in)
        out_kbps = convert_to_kbps(val_out, unit_out)

        total_avg_kbps = in_kbps + out_kbps

        est_monthly_kb = total_avg_kbps * 60 * 60 * 24 * days_in_month

        if not item.get("selected_url"):
            status = "Tidak Ada Grafik"
        elif total_avg_kbps == 0:
            status = "Data Nol / OCR Gagal Membaca Grafik"
        else:
            status = "Sukses"

        rows.append(
            {
                "No": i,
                "Alamat Cabang": item["alamat"],
                "SID": item["sid"],
                "Bandwidth": item["bw"],
                "Avg Inbound (Kbps)": in_kbps,
                "Avg Outbound (Kbps)": out_kbps,
                "Total Average (Kbps)": total_avg_kbps,
                "Estimasi Trafik Bulanan (Kb)": est_monthly_kb,
                "Status": status,
            }
        )

    df = pd.DataFrame(rows)

    safe_name = target_date_str.replace(":", ".").replace("/", "-")
    fname = f"Laporan MRTG BANK JATIM {safe_name}.xlsx"

    df.to_excel(fname, index=False)

    wb = load_workbook(fname)
    ws = wb.active

    center_align = Alignment(horizontal="center", vertical="center")
    for cell in ws[1]:
        cell.alignment = center_align

    left_align = Alignment(horizontal="left", vertical="center")
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = left_align

    yellow_fill = PatternFill(
        start_color="FFE599", end_color="FFE599", fill_type="solid"
    )
    red_fill = PatternFill(start_color="EA9999", end_color="EA9999", fill_type="solid")

    status_col_idx = df.columns.get_loc("Status") + 1

    for r in range(2, ws.max_row + 1):
        cell_status = ws.cell(row=r, column=status_col_idx)
        val = cell_status.value

        fill_color = None
        if val == "Tidak Ada Grafik":
            fill_color = yellow_fill
        elif val == "Data Nan":
            fill_color = red_fill

        if fill_color:
            for c in range(1, ws.max_column + 1):
                ws.cell(row=r, column=c).fill = fill_color

    wb.save(fname)
    return fname


# DOWNLOAD BUTTON
def make_download_button(file_path, label, css_class, mime, icon_rel_path):
    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()

    with open(resource_path(icon_rel_path), "rb") as ic:
        icon_b64 = base64.b64encode(ic.read()).decode()

    filename = os.path.basename(file_path)

    return f"""
    <a class="dl-btn {css_class}"
       href="data:{mime};base64,{b64}"
       download="{filename}">
        <img src="data:image/png;base64,{icon_b64}" class="dl-icon"/>
        <span>{label}</span>
    </a>
    """


# RESET BUTTON
def make_reset_button(label, icon_rel_path):
    try:
        with open(resource_path(icon_rel_path), "rb") as ic:
            icon_b64 = base64.b64encode(ic.read()).decode()
        img_tag = f'<img src="data:image/png;base64,{icon_b64}" class="reset-icon"/>'
    except:
        img_tag = ""

    return f"""
    <a class="reset-btn" href="/?reset=1" target="_self">
        {img_tag}
        <span>{label}</span>
    </a>
    """


# ICON
def img_to_base64(rel_path: str) -> str:
    p = resource_path(rel_path)
    with open(p, "rb") as f:
        return base64.b64encode(f.read()).decode()


def asset_exists(rel_path: str) -> bool:
    return os.path.exists(resource_path(rel_path))


ARROW_ICON = img_to_base64("assets/arrow.png")

# MAIN PROCESS
if asset_exists("assets/telkom.jpg"):
    st.image(resource_path("assets/telkom.jpg"), width=200)
else:
    st.markdown("### Telkom Indonesia")

params = st.query_params
if params.get("reset") == ["1"]:
    st.session_state.clear()
    st.query_params.clear()
    st.session_state.step = "input"
    st.rerun()

if "step" not in st.session_state:
    st.session_state.step = "input"

if st.session_state.step == "input":
    try:
        loading_icon_b64 = img_to_base64("assets/loading.png")
    except:
        loading_icon_b64 = ""

    st.markdown(
        f"""
        <style>
            button[kind="primary"] {{
                background-color: #EE2D24 !important; 
                border: 2px solid #EE2D24 !important; 
                color: white !important;
                border-radius: 50px !important; 
                padding: 0.5rem 2rem !important; 
                font-weight: bold !important;
                transition: all 0.2s ease-in-out !important; 
                box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
            }}

            button[kind="primary"]:hover {{
                background-color: #cc0000 !important; 
                border-color: #cc0000 !important; 
                transform: scale(1.02) !important;
            }}

            button[kind="secondary"] {{
                background-color: white !important; 
                border: 2px solid #EE2D24 !important; 
                color: #EE2D24 !important;
                border-radius: 50px !important; 
                padding: 0.5rem 2rem !important; 
                font-weight: bold !important;
                transition: all 0.2s ease-in-out !important; 
                box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
            }}

            button[kind="secondary"]:hover {{
                background-color: #cc0000 !important; 
                border-color: #cc0000 !important; 
                color: white !important; 
                transform: scale(1.02) !important;
            }}

            [data-testid="stFileUploader"] button {{
                display: inline-flex !important; 
                justify-content: center !important; 
                align-items: center !important;
                margin-top: 5px !important; 
                padding: 8px 25px !important; 
                font-weight: normal !important;
                background-color: #EE2D24 !important; 
                border: 2px solid #EE2D24 !important; 
                color: white !important;
                border-radius: 50px !important; 
                transition: all 0.2s ease-in-out !important;
            }}

            [data-testid="stFileUploader"] section {{ 
                background-color: #f8f9fa !important; 
                border-radius: 15px !important; 
                border: 1px dashed #ccc !important; 
            }}

            @keyframes slideIn {{ 
                from {{ 
                    transform: translateX(100%); 
                    opacity: 0; 
                }} 
                to {{ 
                    transform: translateX(0); 
                    opacity: 1; 
                }} 
            }}

            @keyframes spin {{ 
                0% {{ 
                    transform: rotate(0deg); 
                }} 
                100% {{ 
                    transform: rotate(360deg); 
                }} 
            }}
            
            .custom-toast {{
                position: fixed; 
                top: 20px; 
                right: 20px;
                background-color: white; 
                border-left: 5px solid #EE2D24;
                padding: 15px 20px; 
                border-radius: 8px;
                box-shadow: 0 4px 15px rgba(0,0,0,0.15);
                display: flex; 
                align-items: center; 
                gap: 15px;
                z-index: 999999; 
                animation: slideIn 0.3s ease-out forwards;
                min-width: 320px;
            }}

            .toast-icon {{
                width: 30px; 
                height: 30px;
                background-image: url('data:image/png;base64,{loading_icon_b64}');
                background-size: contain; 
                background-repeat: no-repeat; 
                background-position: center;
                animation: spin 1s linear infinite;
            }}

            .toast-content {{ 
                font-family: 'Segoe UI', sans-serif; 
                flex-grow: 1; 
            }}

            .toast-title {{ 
                font-weight: 800; 
                font-size: 14px; 
                color: #333; 
                margin-bottom: 2px; 
            }}

            .toast-desc {{ 
                font-size: 13px; 
                color: #666; 
                font-weight: 500; 
            }}

            .toast-progress {{ 
                font-size: 11px; 
                color: #999; 
                margin-top: 2px; 
            }}
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        '<div class="main-title">Dashboard Monitoring & Otomatisasi Laporan</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="sub-title">Automasi Monitoring, Validasi Grafik, dan Generasi Laporan</div>',
        unsafe_allow_html=True,
    )

    with st.container(border=True):
        st.write("**Upload File Excel**")
        file = st.file_uploader(
            "", type=["xlsx"], label_visibility="collapsed", key="file_uploader"
        )
        st.markdown(
            '<div class="upload-note">⚠️ Excel harus memiliki kolom Alamat, SID, dan Bandwidth</div>',
            unsafe_allow_html=True,
        )

        st.write("**Konfigurasi Monitoring**")

        col_start, col_end = st.columns(2)
        with col_start:
            st.caption("Mulai")
            s_date = st.date_input(
                "Tanggal Mulai", value=datetime.now().date(), key="s_date"
            )
            s_time_str = st.text_input(
                "Jam Mulai",
                value="00:00",
                key="s_time_str",
                help="Format 24 Jam, contoh: 07:00 atau 23:59",
            )
            try:
                s_time = datetime.strptime(s_time_str, "%H:%M").time()
            except ValueError:
                st.error("Format Jam Mulai salah! Gunakan HH:MM (Cth: 07:00)")
                s_time = datetime.strptime("00:00", "%H:%M").time()

        with col_end:
            st.caption("Berakhir")
            e_date = st.date_input(
                "Tanggal Berakhir", value=datetime.now().date(), key="e_date"
            )
            e_time_str = st.text_input(
                "Jam Berakhir",
                value="23:59",
                key="e_time_str",
                help="Format 24 Jam, contoh: 07:00 atau 23:59",
            )
            try:
                e_time = datetime.strptime(e_time_str, "%H:%M").time()
            except ValueError:
                st.error("Format Jam Berakhir salah! Gunakan HH:MM (Cth: 21:00)")
                e_time = datetime.strptime("23:59", "%H:%M").time()

        st.markdown("<br>", unsafe_allow_html=True)

        st.write("**Tipe Grafik**")
        graph_type = st.selectbox(
            "Pilih Tipe Grafik",
            options=[
                "Daily",
                "Weekly",
                "Monthly",
                "Yearly",
            ],
            index=0,
            label_visibility="collapsed",
        )

        rra_map = {
            "Daily": 1,
            "Weekly": 2,
            "Monthly": 3,
            "Yearly": 4,
        }
        selected_rra_id = rra_map[graph_type]

        st.markdown("<br>", unsafe_allow_html=True)
        btn_proc, btn_cancel, _ = st.columns([1.5, 1.5, 8])

        with btn_proc:
            is_process = st.button("Process", type="primary", use_container_width=True)
        with btn_cancel:
            is_cancel = st.button("Cancel", type="secondary", use_container_width=True)

        if is_cancel:
            st.session_state.clear()
            st.rerun()

        if is_process:
            if file:
                progress_toast = st.empty()
                df = pd.read_excel(file)
                df.columns = [str(c).strip().upper() for c in df.columns]

                final_res = []
                total_data = len(df)

                str_start = f"{s_date.strftime('%Y-%m-%d')} {s_time.strftime('%H:%M')}"
                str_end = f"{e_date.strftime('%Y-%m-%d')} {e_time.strftime('%H:%M')}"

                bulan_indo = [
                    "Januari",
                    "Februari",
                    "Maret",
                    "April",
                    "Mei",
                    "Juni",
                    "Juli",
                    "Agustus",
                    "September",
                    "Oktober",
                    "November",
                    "Desember",
                ]

                def fmt_indo(d):
                    return f"{d.day} {bulan_indo[d.month - 1]} {d.year}"

                if graph_type == "Daily":
                    display_range = fmt_indo(s_date)
                else:
                    display_range = f"{fmt_indo(s_date)} sd {fmt_indo(e_date)}"

                for idx, row in df.iterrows():
                    current_num = idx + 1
                    percent = int((current_num / total_data) * 100)
                    sid_now = str(row["SID"])

                    progress_toast.markdown(
                        f"""
                        <div class="custom-toast">
                            <div class="toast-icon"></div>
                            <div class="toast-content">
                                <div class="toast-title">Memproses SID: {sid_now}</div>
                                <div class="toast-desc">Mode: {graph_type}</div>
                                <div class="toast-progress">{percent}% selesai</div>
                            </div>
                        </div>
                    """,
                        unsafe_allow_html=True,
                    )

                    graphs = scrape_dynamic(
                        sid_now, str_start, str_end, selected_rra_id
                    )

                    final_res.append(
                        {
                            "alamat": row["ALAMAT"],
                            "sid": row["SID"],
                            "bw": row["BANDWIDTH"],
                            "tanggal": display_range,
                            "graphs": graphs,
                            "selected_url": (
                                graphs[0]["url"] if len(graphs) == 1 else None
                            ),
                        }
                    )

                st.session_state.update(
                    {
                        "results": final_res,
                        "target_date_obj": s_date,
                        "display_range_str": display_range,
                        "step": "validate",
                        "current_page": 1,
                        "q_validate": "",
                        "f_validate": "Semua",
                    }
                )
                st.rerun()
            else:
                st.error("Silakan upload file Excel terlebih dahulu.")

elif st.session_state.step == "validate":
    try:
        add_icon_b64 = img_to_base64("assets/add.png")
    except:
        add_icon_b64 = ""

    try:
        loading_icon_b64 = img_to_base64("assets/loading.png")
    except:
        loading_icon_b64 = ""

    st.markdown(
        f"""
        <style>
            button[kind="primary"] {{
                background-color: #EE2D24 !important;
                border: 2px solid #EE2D24 !important;
                color: white !important;
                border-radius: 50px !important;
                padding: 0.6rem 1rem !important;
                font-weight: 800 !important;
                box-shadow: 0 4px 6px rgba(238, 45, 36, 0.2) !important;
                transition: all 0.2s ease-in-out !important;
            }}

            button[kind="primary"]:hover {{
                background-color: #cc0000 !important;
                border-color: #cc0000 !important;
                transform: scale(1.02) !important;
            }}

            button[kind="primary"]::before {{
                content: "";
                display: inline-block;
                width: 22px;
                height: 22px;
                background-image: url('data:image/png;base64,{add_icon_b64}');
                background-size: contain;
                background-repeat: no-repeat;
                background-position: center;
                filter: brightness(0) invert(1);
                flex-shrink: 0;
                margin-right: 8px;
            }}

            button[kind="secondary"] {{
                background-color: white !important;
                border: 2px solid #000 !important;
                color: #000 !important;
                border-radius: 12px !important;
                padding: 8px 16px !important;
                font-weight: 600 !important;
                transition: all 0.2s ease !important;
                height: auto !important;
            }}

            button[kind="secondary"]:hover {{
                border-color: #EE2D24 !important;
                color: #EE2D24 !important;
                background-color: #fff5f5 !important;
            }}
            
            button[kind="secondary"]:disabled {{
                border-color: #e5e7eb !important;    
                background-color: #f9fafb !important; 
                color: #9ca3af !important;            
                opacity: 0.7 !important;              
                cursor: not-allowed !important;
            }}

            div[data-testid="column"]:nth-of-type(2) button {{
                border: 2px solid #EE2D24 !important; 
                color: #EE2D24 !important;            
                background-color: white !important;
                border-radius: 10px !important;       
                padding: 0px !important;
                height: 46px !important;
                width: 100% !important;
                display: flex !important; 
                align-items: center !important; 
                justify-content: center !important;
            }}

            div[data-testid="column"]:nth-of-type(2) button:hover {{
                background-color: #EE2D24 !important; 
                color: white !important;              
                border-color: #EE2D24 !important;
                box-shadow: 0 4px 8px rgba(238, 45, 36, 0.3) !important;
                transform: translateY(-2px);
            }}

            div[data-testid="column"]:nth-of-type(2) button p {{
                font-size: 28px !important;
                font-weight: 900 !important;
                line-height: 1 !important;
                margin-top: -5px !important;
            }}

            @keyframes slideIn {{
                from {{
                    transform: translateX(100%);
                    opacity: 0;
                }}
                to {{
                    transform: translateX(0);
                    opacity: 1;
                }}
            }}

            @keyframes spin {{
                0% {{
                    transform: rotate(0deg);
                }}
                100% {{
                    transform: rotate(360deg);
                }}
            }}
            
            .custom-toast {{
                position: fixed;
                top: 20px;
                right: 20px;
                background-color: white;
                border-left: 5px solid #EE2D24;
                padding: 15px 20px;
                border-radius: 8px;
                box-shadow: 0 4px 15px rgba(0,0,0,0.15);
                display: flex;
                align-items: center;
                gap: 15px;
                z-index: 999999;
                animation: slideIn 0.5s ease-out forwards;
                min-width: 300px;
            }}

            .toast-icon {{
                width: 30px;
                height: 30px;
                background-image: url('data:image/png;base64,{loading_icon_b64}');
                background-size: contain;
                background-repeat: no-repeat;
                background-position: center;
                animation: spin 1s linear infinite; 
            }}

            .toast-content {{
                font-family: 'Segoe UI', sans-serif;
            }}

            .toast-title {{
                font-weight: 800;
                font-size: 14px;
                color: #333;
                margin-bottom: 2px;
            }}

            .toast-desc {{
                font-size: 12px;
                color: #666;
            }}

            .pag-center {{
                display: flex;
                justify-content: center;
                align-items: center;
                height: 100%;
                font-size: 16px;
                color: #444;
                padding-top: 10px;
            }}
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="main-title">Validasi Grafik</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="sub-title">Jika ditemukan lebih dari satu grafik, silakan pilih yang paling akurat.</div>',
        unsafe_allow_html=True,
    )

    if "current_page" not in st.session_state:
        st.session_state.current_page = 1
    if "q_validate" not in st.session_state:
        st.session_state.q_validate = ""
    if "f_validate" not in st.session_state:
        st.session_state.f_validate = "Semua"

    def reset_callback():
        st.session_state.q_validate = ""
        st.session_state.f_validate = "Semua"
        st.session_state.current_page = 1

    c_search, c_reset, c_filter, c_btn = st.columns([5, 1, 3, 3], gap="small")

    with c_search:
        st.text_input(
            "Cari Cabang",
            placeholder="Cari Nama Alamat atau SID",
            label_visibility="collapsed",
            key="q_validate",
        )

    with c_reset:
        st.button(
            "↺",
            key="btn_reset_text",
            help="Reset Filter & Pencarian",
            on_click=reset_callback,
            use_container_width=True,
            type="secondary",
        )

    with c_filter:
        st.selectbox(
            "Filter Status",
            options=["Semua", "Berhasil", "Perlu Validasi", "Tidak Ada Grafik"],
            label_visibility="collapsed",
            key="f_validate",
        )

    with c_btn:
        do_print = st.button(
            "CETAK SEMUA LAPORAN", type="primary", use_container_width=True
        )

    st.markdown("<div style='margin-bottom: 20px;'></div>", unsafe_allow_html=True)

    total_data = len(st.session_state.results)
    sukses = sum(1 for x in st.session_state.results if len(x.get("graphs", [])) == 1)
    perlu_cek = sum(1 for x in st.session_state.results if len(x.get("graphs", [])) > 1)
    gagal = sum(1 for x in st.session_state.results if len(x.get("graphs", [])) == 0)

    m1, m2, m3, m4 = st.columns(4)
    with m1:
        st.markdown(
            f'<div class="metric-box mb-blue"><div class="metric-title">TOTAL CABANG</div><div class="metric-value">{total_data}</div></div>',
            unsafe_allow_html=True,
        )
    with m2:
        st.markdown(
            f'<div class="metric-box mb-green"><div class="metric-title">✅ BERHASIL</div><div class="metric-value">{sukses}</div></div>',
            unsafe_allow_html=True,
        )
    with m3:
        st.markdown(
            f'<div class="metric-box mb-yellow"><div class="metric-title">⚠️ PERLU VALIDASI</div><div class="metric-value">{perlu_cek}</div></div>',
            unsafe_allow_html=True,
        )
    with m4:
        st.markdown(
            f'<div class="metric-box mb-red"><div class="metric-title">❌ TIDAK ADA GRAFIK</div><div class="metric-value">{gagal}</div></div>',
            unsafe_allow_html=True,
        )

    st.markdown("<br>", unsafe_allow_html=True)

    def compute_status(cnt):
        if cnt == 0:
            return "Tidak Ada Grafik"
        if cnt == 1:
            return "Berhasil"
        return "Perlu Validasi"

    query = (st.session_state.q_validate or "").strip()
    f_stat = st.session_state.f_validate

    processed = []
    for i, item in enumerate(st.session_state.results):
        cnt = len(item.get("graphs", []))
        status = compute_status(cnt)
        q_ok = (
            (query == "")
            or (query.lower() in item["alamat"].lower())
            or (query in str(item["sid"]))
        )
        f_ok = (f_stat == "Semua") or (f_stat == status)
        if q_ok and f_ok:
            it = dict(item)
            it["orig_idx"] = i
            it["status"] = status
            processed.append(it)

    items_per_page = 5
    total_pages = max(1, (len(processed) + items_per_page - 1) // items_per_page)
    st.session_state.current_page = max(
        1, min(st.session_state.current_page, total_pages)
    )
    start_idx = (st.session_state.current_page - 1) * items_per_page
    paged_data = processed[start_idx : start_idx + items_per_page]

    for item in paged_data:
        idx = item["orig_idx"]
        status = item["status"]
        card_class = (
            "card-success"
            if status == "Berhasil"
            else "card-failed" if status == "Tidak Ada Grafik" else "card-multiple"
        )
        badge = (
            "✅ Berhasil"
            if status == "Berhasil"
            else (
                "❌ Tidak Ada Grafik"
                if status == "Tidak Ada Grafik"
                else "⚠️ Perlu Validasi"
            )
        )

        st.markdown(
            f"""
        <div class="{card_class}">
            <div class="card-title">{idx+1}. {item['alamat']}</div>
            <div class="card-meta">SID: {item['sid']} &nbsp;|&nbsp; BW: {item['bw']} &nbsp;|&nbsp; <b>{badge}</b></div>
        </div>
        """,
            unsafe_allow_html=True,
        )

        with st.container():
            if status == "Tidak Ada Grafik":
                st.error("Tidak Ditemukan Grafik.")
            elif status == "Berhasil":
                st.image(item["graphs"][0]["url"], width=520)
                st.session_state.results[idx]["selected_url"] = item["graphs"][0]["url"]
            else:
                graphs = item["graphs"]
                if (
                    not st.session_state.results[idx].get("selected_url")
                    and len(graphs) > 0
                ):
                    st.session_state.results[idx]["selected_url"] = graphs[0]["url"]
                current_selected = st.session_state.results[idx].get("selected_url")

                for row_start in range(0, len(graphs), 3):
                    row_graphs = graphs[row_start : row_start + 3]
                    cols = st.columns(3)
                    for col_i, g in enumerate(row_graphs):
                        with cols[col_i]:
                            st.image(g["url"], use_container_width=True)
                            is_selected = current_selected == g["url"]
                            label = (
                                "☑ Terpilih"
                                if is_selected
                                else f"Opsi {row_start + col_i + 1}"
                            )

                            if st.button(
                                label,
                                key=f"pick_{idx}_{row_start + col_i}",
                                type="secondary",
                                use_container_width=True,
                                disabled=is_selected,
                            ):
                                st.session_state.results[idx]["selected_url"] = g["url"]
                                st.rerun()
                    st.markdown(
                        "<div style='height:10px'></div>", unsafe_allow_html=True
                    )
        st.divider()

    if do_print:
        st.markdown(
            f"""
            <div class="custom-toast">
                <div class="toast-icon"></div>
                <div class="toast-content">
                    <div class="toast-title">Sedang Memproses..</div>
                    <div class="toast-desc">Membaca Grafik & Membuat Laporan</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        for i, it in enumerate(st.session_state.results):
            if it.get("selected_url"):
                try:
                    img_data = requests.get(it["selected_url"], timeout=10).content
                    tmp_path = f"tmp_ocr_{i}.png"
                    with open(tmp_path, "wb") as f:
                        f.write(img_data)

                    a_in, a_out, m_in, m_out = ocr_extract_data(tmp_path)
                    os.remove(tmp_path)

                    st.session_state.results[i].update(
                        {
                            "avg_in": a_in,
                            "avg_out": a_out,
                            "max_in": m_in,
                            "max_out": m_out,
                        }
                    )
                except Exception as e:
                    st.session_state.results[i].update(
                        {"avg_in": "0", "avg_out": "0", "max_in": "0", "max_out": "0"}
                    )

        tgl_display = st.session_state.get("display_range_str", "Laporan Monitoring")

        safe_filename_date = tgl_display.replace(":", ".").replace("/", "-")

        path_word = generate_clean_word(st.session_state.results, tgl_display)

        tgl_objek_asli = st.session_state.get("target_date_obj", None)

        path_excel = generate_excel_report(
            st.session_state.results,
            safe_filename_date,
            date_obj=tgl_objek_asli,
        )

        st.session_state.update(
            {"final_path": path_word, "final_excel": path_excel, "step": "finish"}
        )
        st.rerun()

    c_prev, c_mid, c_next = st.columns([3, 4, 3])
    with c_prev:
        if st.button(
            "« Previous",
            type="secondary",
            use_container_width=True,
            disabled=(st.session_state.current_page == 1),
            key="pg_prev",
        ):
            st.session_state.current_page -= 1
            st.rerun()
    with c_mid:
        st.markdown(
            f"<div class='pag-center'>Halaman&nbsp;<b>{st.session_state.current_page}</b>&nbsp;dari&nbsp;<b>{total_pages}</b></div>",
            unsafe_allow_html=True,
        )
    with c_next:
        if st.button(
            "Next »",
            type="secondary",
            use_container_width=True,
            disabled=(st.session_state.current_page == total_pages),
            key="pg_next",
        ):
            st.session_state.current_page += 1
            st.rerun()

elif st.session_state.step == "finish":
    st.markdown(
        """
        <style>
            .finish-title {
                font-size: 36px;
                font-weight: 800;
                color: #000;
                text-align: center;
                margin-bottom: 20px;
                font-family: 'Segoe UI', sans-serif;
            }

            .dl-btn, .reset-btn {
                display: flex;
                align-items: center;
                justify-content: center;
                gap: 12px;
                height: 60px; 
                border-radius: 999px;
                font-weight: 800;
                font-size: 18px;
                text-decoration: none !important;
                border: 3px solid #000;
                box-shadow: 0 4px 10px rgba(0,0,0,0.1);
                transition: all 0.2s ease;
                width: 100%;
                cursor: pointer;
            }

            .dl-btn:hover, .reset-btn:hover {
                transform: translateY(-2px);
                box-shadow: 0 6px 15px rgba(0,0,0,0.15);
            }

            .btn-blue {
                background: #2D62D3;
                color: #fff !important;
                border-color: #2D62D3 !important;
            }

            .btn-green {
                background: #1E7A3A;
                color: #fff !important;
                border-color: #1E7A3A !important;
            }

            .reset-btn {
                background: #ffffff;
                color: #000 !important;
                border-color: #000 !important;
            }

            .dl-icon, .reset-icon {
                width: 24px;
                height: 24px;
                object-fit: contain;
            }

            .download-row {
                margin-top: 40px;
            }

            .reset-row {
                margin-top: 20px;
            }

            @keyframes bounce-check {
                0%, 100% {
                    transform: translateY(0);
                }
                40% {
                    transform: translateY(-15px);
                }
                60% {
                    transform: translateY(-7px);
                }
            }

            .check-bounce {
                animation: bounce-check 1.8s ease-in-out infinite;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div style="display:flex;justify-content:center;margin-bottom:20px;margin-top:20px;">
            <div class="check-bounce" style="
                width:100px;height:100px;background:#EE2D24;border-radius:50%;
                display:flex;align-items:center;justify-content:center;
                box-shadow: 0 10px 25px rgba(238, 45, 36, 0.3);">
                <svg width="50" height="50" viewBox="0 0 24 24" fill="none" stroke="white" stroke-width="4" stroke-linecap="round" stroke-linejoin="round">
                    <polyline points="20 6 9 17 4 12"></polyline>
                </svg>
            </div>
        </div>
    """,
        unsafe_allow_html=True,
    )

    st.markdown(
        '<div class="finish-title">Laporan Berhasil Dibuat!</div>',
        unsafe_allow_html=True,
    )

    icon_dl = (
        "assets/downloads.png"
        if asset_exists("assets/downloads.png")
        else "assets/arrow.png"
    )

    html_word = make_download_button(
        st.session_state.final_path,
        "Download Word",
        "btn-blue",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        icon_dl,
    )
    html_excel = make_download_button(
        st.session_state.final_excel,
        "Download Excel",
        "btn-green",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        icon_dl,
    )

    st.markdown('<div class="download-row">', unsafe_allow_html=True)
    cL, cWord, cExcel, cR = st.columns([2, 3, 3, 2], gap="large")
    with cWord:
        st.markdown(html_word, unsafe_allow_html=True)
    with cExcel:
        st.markdown(html_excel, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="reset-row">', unsafe_allow_html=True)
    rL, rMid, rR = st.columns([3, 4, 3])
    with rMid:
        icon_reset = (
            "assets/reset.png"
            if asset_exists("assets/reset.png")
            else "assets/arrow.png"
        )
        html_reset = make_reset_button("Buat Laporan Baru", icon_reset)
        st.markdown(html_reset, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)
